from openai import OpenAI
from docx import Document

# Set up OpenAI API key
client = OpenAI(api_key="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")

model = "gpt-4o-mini"
PRICING_PER_1M_TOKENS = {
    "input": 0.150,  # $0.150 per 1M input tokens
    "output": 0.600  # $0.600 per 1M output tokens
}
OUTPUT_TOKEN_FACTOR = 0.5  # Assume output tokens are 50% of input tokens

def parse_word_document(file_path):
    """Extract text from a Word document."""
    doc = Document(file_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

def parse_and_estimate_tokens_cost(file_path):
    """Parse the Word document and estimate its token count."""
    doc = Document(file_path)
    document_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    estimated_tokens = estimate_tokens(document_text)
    cost = calculate_cost(estimated_tokens, estimated_tokens*OUTPUT_TOKEN_FACTOR)
    print(f"Estimated tokens in the document: {estimated_tokens}")
    print(f"Estimated cost USD: {cost}")
    return document_text, estimated_tokens

def calculate_cost(input_tokens, output_tokens):
    """Calculate the cost of processing the given number of tokens."""
    input_cost = (input_tokens / 1_000_000) * PRICING_PER_1M_TOKENS["input"]
    output_cost = (output_tokens / 1_000_000) * PRICING_PER_1M_TOKENS["output"]
    return input_cost + output_cost

def estimate_tokens(text):
    """Estimate the number of tokens in a text based on character count."""
    avg_chars_per_token = 4  # Average characters per token
    return len(text) // avg_chars_per_token

def summarize_resume(document_text):
    """Use an LLM to summarize the resume into relevant chunks."""
    print("Summarizing the document using LLM...")
    prompt = f"""You are a helpful assistant that extracts and summarizes resumes.
    Given the following resume text, create a concise summary of relevant sections such as 'Experience', 'Education', 'Skills', 'Projects', etc. Use as few words as possible without losing meaning.
    Resume Text:
    {document_text}
    Provide the output as a JSON object where section titles are keys and their summaries are values."""
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000,  # Adjust if necessary based on expected output size
        temperature=0.3
    )
    return response.choices[0].message.content

def process_resume(file_path, output_file):
    """Process the resume and save summarized contexts to a file."""
    document_text = parse_word_document(file_path)
    estimated_tokens = estimate_tokens(document_text)
    cost = calculate_cost(estimated_tokens, estimated_tokens*OUTPUT_TOKEN_FACTOR)
    print(f"Estimated tokens in the document: {estimated_tokens}")
    print(f"Estimated cost USD: {cost} for model {model}")
    confirm = input("Do you want to continue Y/N: ").strip()
    if confirm.lower() == 'n':
        print("exiting")
        return
    summary = summarize_resume(document_text)
    with open(output_file, "w") as f:
        f.write(summary)
    print(f"Summarized resume saved to {output_file}")

def main():
    """Main function to process a Word resume and save summaries."""
    file_path = input("Enter the path to the Word document resume: ").strip()
    output_file = input("Enter the path to save the summarized contexts (JSON file): ").strip()
    process_resume(file_path, output_file)

if __name__ == "__main__":
    main()
